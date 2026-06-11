import csv
import re
from dataclasses import dataclass
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import CHUNK_SIZE

CSV_ENCODINGS = ("utf-8-sig", "utf-8", "latin-1", "cp1252", "utf-16")


@dataclass
class ExtractedDocument:
    text: str
    row_count: int | None = None


def _decode_csv(content: bytes) -> str:
    for encoding in CSV_ENCODINGS:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _detect_delimiter(sample: str) -> str:
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        return dialect.delimiter
    except csv.Error:
        if ";" in sample.splitlines()[0]:
            return ";"
        if "\t" in sample.splitlines()[0]:
            return "\t"
        return ","


def _extract_csv(content: bytes) -> ExtractedDocument:
    text = _decode_csv(content)
    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return ExtractedDocument(text="")

    delimiter = _detect_delimiter("\n".join(lines[:5]))
    reader = csv.reader(StringIO("\n".join(lines)), delimiter=delimiter)
    parsed = [[str(cell).strip() for cell in row] for row in reader]
    parsed = [row for row in parsed if any(cell for cell in row)]
    if not parsed:
        return ExtractedDocument(text="")

    header = parsed[0]
    has_header = len(parsed) > 1 and all(header)
    data_rows = parsed[1:] if has_header else parsed

    rows: list[str] = []
    if has_header:
        rows.append("Columns: " + ", ".join(header))
        for row in data_rows:
            pairs = []
            for col, val in zip(header, row):
                if val:
                    pairs.append(f"{col}={val}")
            if pairs:
                rows.append(", ".join(pairs))
    else:
        for row in parsed:
            rows.append(" | ".join(cell for cell in row if cell))

    return ExtractedDocument(text="\n".join(rows), row_count=len(data_rows))


def _extract_docx(content: bytes) -> ExtractedDocument:
    doc = Document(BytesIO(content))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return ExtractedDocument(text="\n\n".join(paragraphs))


def _extract_xlsx(content: bytes) -> ExtractedDocument:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    total_rows = 0
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            total_rows += max(0, len(rows) - 1)
            parts.append(f"Sheet: {sheet.title}\n" + "\n".join(rows))
    workbook.close()
    return ExtractedDocument(text="\n\n".join(parts), row_count=total_rows or None)


def extract_document(filename: str, content: bytes, content_type: str = "") -> ExtractedDocument:
    suffix = Path(filename).suffix.lower() if filename else ""
    if not suffix and "csv" in content_type.lower():
        suffix = ".csv"

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return ExtractedDocument(text="\n\n".join(pages).strip())

    if suffix in {".txt", ".md"}:
        return ExtractedDocument(text=content.decode("utf-8", errors="replace").strip())

    if suffix == ".csv":
        return _extract_csv(content)

    if suffix == ".docx":
        return _extract_docx(content)

    if suffix == ".xlsx":
        return _extract_xlsx(content)

    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")


def chunk_text(text: str, is_csv: bool = False) -> list[str]:
    if is_csv:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return []

        header = lines[0] if lines[0].startswith("Columns:") else None
        data_lines = lines[1:] if header else lines

        chunks: list[str] = []
        batch: list[str] = []
        batch_len = 0

        for line in data_lines:
            line_len = len(line) + 1
            if batch and batch_len + line_len > CHUNK_SIZE:
                chunk_lines = ([header] if header else []) + batch
                chunks.append("\n".join(chunk_lines))
                batch = [line]
                batch_len = line_len
            else:
                batch.append(line)
                batch_len += line_len

        if batch:
            chunk_lines = ([header] if header else []) + batch
            chunks.append("\n".join(chunk_lines))

        return chunks

    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []

    chunks = []
    start = 0
    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - 150
    return chunks
